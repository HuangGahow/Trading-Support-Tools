# -*- coding: utf-8 -*-
"""
Created on Sat Nov  5 23:08:35 2022

@author: HP
"""

import docx
import pandas as pd
import numpy as np
from datetime import datetime
import os

# In[1]


def DateStrtoTime(x):
    if type(x) == datetime:
        stamp = x
        
    else:
        test = x.split(',')
    
        calendar={'January':'01','February':'02','March':'03','April':'04', \
                  'May':'05','June':'06','July':'07','August':'08','September':'09', \
                      'October':'10','November':'11', 'December':'12'}
        day = int(test[0])                
        month = int(calendar[test[1]])
        year = int(test[2])
    
        stamp = datetime(year,month,day)
    
    return stamp


def GetEquityAmounttext(x):
    EAindex1  = x.find('(')  #find position of (
    #EAindex2  = x.find(' -Commission')
    EquityAmount = x[EAindex1+1]
    return EquityAmount


def changevalue(x):
    if type(x) == float:
        return x
    else:
        a = x.split(' ')[1] #seperate CNY/USD
        a = ''.join(a.split(',')) #seperate ',' and combine numbers
        result = float(a)
        return result

# In[3]

def Get_Confirm_Data(fn):

    doc = docx.Document(fn)
    pageheader = doc.sections[0].header.paragraphs[0].text #page header
    
    #table_0 = doc.tables[0] #get first table in word
    date = doc.tables[0].rows[0].cells[1].text.strip() #get 0 row 1st cell in 1st table
    contract = doc.tables[0].rows[3].cells[1].text
    
    
    

    TradeDate, EffectiveDate = doc.tables[1].rows[4].cells[1].text, \
                                            doc.tables[1].rows[5].cells[1].text
                                            #EffectiveDate 起始日，交易日                                              
    ValuationDate  = doc.tables[1].rows[6].cells[1].text.split(',')[0:3]   
    ValuationDate  = ','.join(ValuationDate)
    date,ValuationDate, TradeDate, EffectiveDate  =DateStrtoTime(date), DateStrtoTime(ValuationDate), \
        DateStrtoTime(TradeDate),DateStrtoTime(EffectiveDate)
        
        
    TradeReferenceNumber = doc.tables[1].rows[12].cells[1].text 


    A = doc.tables[1].rows[14].cells[1].text + doc.tables[1].rows[15].cells[1].text
    A = A.strip()
    FuturesContract = A.split('(')[1].split(')')[0]
    Contractname = A.split('(')[0]
    contactdict={'5-year':'F','10-year':'2','2-year':'S'}
    #check if contact number is matched with name 
    ContractnameInterVeri = (contactdict[Contractname.split()[0]] ==  FuturesContract[1])    
    
    

        

    EquityAmount = doc.tables[2].rows[3].cells[1].text
    EquityAmount = GetEquityAmounttext(EquityAmount)
    if EquityAmount =='F': #(Final price - initial price)
        BorS = 'B'
    elif EquityAmount =='I':
        BorS = 'S'
    else:
        BorS = np.nan



    NotionalAmount,InitialPrice = doc.tables[1].rows[17].cells[1].text, \
        doc.tables[1].rows[18].cells[1].text 
    
    NotionalAmount = changevalue(NotionalAmount)
    InitialPrice =  changevalue(InitialPrice)

    ContractNumberInterVerified = (pageheader==contract==TradeReferenceNumber)
    TEDateInterVerified = (TradeDate==EffectiveDate==date)



    d1 = np.array([TradeReferenceNumber,date,TradeDate, EffectiveDate, ValuationDate,FuturesContract, \
                   NotionalAmount,InitialPrice, EquityAmount,BorS,ContractNumberInterVerified, \
                       TEDateInterVerified,ContractnameInterVeri]).reshape(1,-1)
    
    columnnames = ['交易确认书编号','Date','TradeDate', 'EffectiveDate', 'ValuationDate','FuturesContract', \
                   'NotionalAmount','InitialPrice', 'EquityAmount','BorS','InterNumberVeri','InterDateVeri', \
                       'InterContractnameVeri'] 
    
    d1 = pd.DataFrame(d1,columns=columnnames)
    
    return d1


                
# In[4]
path = input("please input your 北向国债期货连接端's path:")
allKeyPath = input("please input your file folder's path:")
bond = pd.read_excel(path, sheet_name="Bond")
#delete the rows include missing value
bond = bond.dropna(axis=0, subset=['交易确认书编号'])
Date = bond['起始日'].iloc[-1]
bond = bond.loc[bond['起始日'] == Date]
bond.drop(["协议", "方向" ,"客户名称","状态","系统编号","交易达成日", "结算金额（客户方向）", \
           "新备案账号", "标的小类" ,"备注", "Unnamed: 28" ],axis = 1, inplace = True)
    


files = os.listdir(allKeyPath)
file_path = []

for file in files:
   # combine file's path
    file_path.append(allKeyPath + '\\' + file)
    


d0 = pd.DataFrame() #d0 is data from Excel

for filename in file_path:
    
    fn =  filename
    d1 = Get_Confirm_Data(fn)
    d0 =d0.append(d1)
# In[5]    

result = bond.merge(d0, left_on='交易确认书编号', right_on='交易确认书编号',how='left')
result['TradeDateVeri'] =(result['起始日']== result['Date'])
result['P0Veri'] =(result['期初价格']== result['InitialPrice'])
result['B/SVeri'] =(result['B/S']== result['BorS'])
result['ValuationDateVeri'] =(result['ValuationDate']== result['到期日'])
result['ContractVeri']=(result['FuturesContract']==result['标的'])
result['NotionalAmountVeri']=(result['NotionalAmount']==result['名义本金USD/总面值'])

report = result.loc[:,['交易确认书编号','TradeDateVeri', \
                       'P0Veri','B/SVeri','ValuationDateVeri', 'NotionalAmountVeri',\
                           'ContractVeri','InterNumberVeri','InterDateVeri', 'InterContractnameVeri']]
    
folder = os.getcwd() + '\\Reports\\'
if not os.path.exists(folder):
    os.makedirs(folder)
report.to_excel(folder + 'Confirmation_Report.xlsx',sheet_name='Report',index = False,header=True)
print("Done")

